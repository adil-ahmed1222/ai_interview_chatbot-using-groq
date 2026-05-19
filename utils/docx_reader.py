"""Extract text from DOCX resumes using python-docx."""

import logging
from io import BytesIO
from typing import BinaryIO, Union

from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file: Union[str, BinaryIO, BytesIO]) -> str:
    """
    Extract plain text from a DOCX file path or file-like object.

    Args:
        file: Path string or binary stream.

    Returns:
        Full document text.

    Raises:
        ValueError: If document is empty.
    """
    try:
        doc = Document(file)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        full_text = "\n".join(paragraphs).strip()
        if not full_text:
            raise ValueError("No readable text found in DOCX file.")
        logger.info("Extracted %d characters from DOCX.", len(full_text))
        return full_text
    except ValueError:
        raise
    except Exception as exc:
        logger.exception("DOCX extraction failed.")
        raise RuntimeError(f"Failed to read DOCX: {exc}") from exc
